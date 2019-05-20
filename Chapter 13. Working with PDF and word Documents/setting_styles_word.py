import docx

doc = docx.Document('demo.docx')
print(doc.paragraphs[0].text)  # 'Document Title'
print(doc.paragraphs[0].style)  # 'Title'
doc.paragraphs[0].style = 'Normal'
print(doc.paragraphs[1].text)  # 'A plain paragraph with some bold and some italic'
print((doc.paragraphs[1].runs[0].text, doc.paragraphs[1].runs[1].text, doc.paragraphs[1].runs[2].text,
       doc.paragraphs[1].runs[3].text))  # ('A plain paragraph with some ', 'bold', ' and some ', 'italic')
doc.paragraphs[1].runs[0].style = 'QuoteChar'
doc.paragraphs[1].runs[1].underline = True
doc.paragraphs[1].runs[3].underline = True
doc.save('restyled.docx')