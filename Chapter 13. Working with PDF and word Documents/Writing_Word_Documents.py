import docx
doc = docx.Document()
doc.add_paragraph('Hello world!')  # <docx.text.Paragraph object at 0x000000000366AD30>
paraObj1 = doc.add_paragraph('This is a second paragraph.')
paraObj2 = doc.add_paragraph('This is a yet another paragraph.')
paraObj1.add_run(' This text is being added to the second paragraph.')  # <docx.text.Run object at 0x0000000003A2C860>
doc.save('multipleParagraphs.docx')
