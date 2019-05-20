# ! python3
# invitations.py - generates custom invitations

import docx
import docx.enum.text


def custom_inv(txt_file):
    with open(txt_file) as f:
        content = f.read().splitlines()

    doc = docx.Document()

    inv_text = ['It would be a pleasure to have the company of', 'at 11010 Memory Lane on the evening of', 'April 1st',
                'at 7 o clock']

    inc = len(inv_text)
    for person in content:
        doc.add_paragraph(inv_text[0], style='Intense Quote')
        doc.add_paragraph(person, style='Subtitle')
        doc.add_paragraph(inv_text[1], style='Subtitle')
        doc.add_paragraph(inv_text[2], style='Subtitle')
        doc.add_paragraph(inv_text[3], style='Subtitle')
        doc.paragraphs[inc].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
        inc += 5
    doc.save('invitations.docx')


print(custom_inv('guests.txt'))
